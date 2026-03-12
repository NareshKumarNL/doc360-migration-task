#!/usr/bin/env python3
import html
import json
import re
from pathlib import Path

import requests
from docx import Document


CONFIG_PATH = Path(__file__).with_name("config.json")


def load_config():
    return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))


def extract_runs(paragraph):
    parts = []
    for run in paragraph.runs:
        text = html.escape(run.text)
        if not text:
            continue
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"
        parts.append(text)
    return "".join(parts).strip()


def list_depth(style_name):
    match = re.search(r"(\d+)$", style_name.strip())
    return int(match.group(1)) if match else 1


def convert_docx_to_html(docx_path):
    doc = Document(str(docx_path))
    html_lines = []
    list_stack = []
    title = None

    def close_lists(target_depth=0):
        while len(list_stack) > target_depth:
            html_lines.append(f"</{list_stack.pop()}>")

    for para in doc.paragraphs:
        raw_text = para.text.strip()
        if not raw_text:
            continue

        style = para.style.name or "Normal"
        text = extract_runs(para)
        if title is None and "Heading 1" in style:
            title = raw_text

        if "Bullet" in style or "Number" in style:
            tag = "ul" if "Bullet" in style else "ol"
            depth = list_depth(style)

            while len(list_stack) > depth:
                html_lines.append(f"</{list_stack.pop()}>")
            while len(list_stack) < depth:
                html_lines.append(f"<{tag}>")
                list_stack.append(tag)
            if list_stack and list_stack[-1] != tag:
                html_lines.append(f"</{list_stack.pop()}>")
                html_lines.append(f"<{tag}>")
                list_stack.append(tag)

            html_lines.append(f"<li>{text}</li>")
            continue

        close_lists(0)

        if "Heading 1" in style:
            html_lines.append(f"<h1>{text}</h1>")
        elif "Heading 2" in style:
            html_lines.append(f"<h2>{text}</h2>")
        elif "Heading 3" in style:
            html_lines.append(f"<h3>{text}</h3>")
        elif "Heading 4" in style:
            html_lines.append(f"<h4>{text}</h4>")
        else:
            html_lines.append(f"<p>{text}</p>")

    close_lists(0)

    for table in doc.tables:
        html_lines.append('<table border="1" cellpadding="10" cellspacing="0" style="border-collapse:collapse;width:100%;margin:20px 0;">')
        for row_index, row in enumerate(table.rows):
            html_lines.append("<tr>")
            for cell in row.cells:
                tag = "th" if row_index == 0 else "td"
                html_lines.append(f"<{tag}>{html.escape(cell.text.strip())}</{tag}>")
            html_lines.append("</tr>")
        html_lines.append("</table>")

    article_body = "\n".join(html_lines)
    page_title = title or Path(docx_path).stem
    full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{html.escape(page_title)}</title>
  <style>
    body {{ font-family: Segoe UI, Arial, sans-serif; line-height: 1.6; margin: 40px auto; max-width: 920px; color: #1f2937; }}
    h1 {{ color: #0f4c81; border-bottom: 3px solid #0f4c81; padding-bottom: 10px; }}
    h2 {{ color: #0f4c81; margin-top: 28px; }}
    h3, h4 {{ color: #233044; margin-top: 20px; }}
    p {{ margin: 12px 0; }}
    ul, ol {{ margin: 12px 0; padding-left: 28px; }}
    li {{ margin: 6px 0; }}
    table {{ width: 100%; margin: 20px 0; border-collapse: collapse; }}
    th {{ background: #0f4c81; color: #fff; text-align: left; }}
    td, th {{ border: 1px solid #d1d5db; padding: 10px; vertical-align: top; }}
  </style>
</head>
<body>
{article_body}
</body>
</html>
"""
    return page_title, article_body, full_html


def slugify(value):
    base = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return base or "migrated-document"


def post_json(url, headers, payload, expected_status=200):
    response = requests.post(url, headers=headers, json=payload, timeout=30)
    body = response.json()
    if response.status_code != expected_status or not body.get("success", False):
        raise RuntimeError(f"{response.status_code}: {json.dumps(body)}")
    return body


def main():
    config = load_config()
    docx_path = Path(config["docx_path"])
    output_html = Path(config["output_html"])
    output_meta = Path(config["output_meta"])

    title, article_body, full_html = convert_docx_to_html(docx_path)
    output_html.write_text(full_html, encoding="utf-8")

    headers = {
        "api_token": config["api_key"],
        "Content-Type": "application/json",
    }
    slug = f'{slugify(title)}-{config["slug_suffix"]}'
    create_payload = {
        "title": title,
        "content": article_body,
        "category_id": config["category_id"],
        "project_version_id": config["project_version_id"],
        "order": 0,
        "user_id": config["user_id"],
        "content_type": 1,
        "slug": slug,
    }

    create_result = post_json(
        "https://apihub.document360.io/v2/Articles",
        headers,
        create_payload,
        expected_status=200,
    )
    article = create_result["data"]

    publish_payload = {
        "user_id": config["user_id"],
        "version_number": article.get("latest_version", 1),
        "publish_message": "Imported DOCX content via API",
    }
    publish_result = post_json(
        f'https://apihub.document360.io/v2/Articles/{article["id"]}/{config["lang_code"]}/publish',
        headers,
        publish_payload,
        expected_status=200,
    )

    output_meta.write_text(
        json.dumps(
            {
                "title": title,
                "html_path": str(output_html),
                "article_id": article["id"],
                "article_url": publish_result["url"],
                "create_status": 200,
                "publish_status": 200,
            },
            indent=2,
        ),
        encoding="utf-8",
    )

    print("HTML saved:", output_html)
    print("Create response: 200")
    print("Publish response: 200")
    print("Article ID:", article["id"])
    print("Article URL:", publish_result["url"])


if __name__ == "__main__":
    main()
